from io import BytesIO
from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from pathlib import Path
import unittest
import zipfile

from app.fixing.docx_editor import apply_fixes_to_source_docx
from app.fixing.document_fixer import apply_fixes, build_corrected_docx
from app.orchestration.pipeline import DocumentPipeline
from app.parsing.document_to_schema import parse_docx_to_document
from app.schemas.document import DocumentInput
from app.standards.ingest import StandardIngestor
from app.standards.registry import StandardRegistry
from tests.support.fake_provider import FakeProvider


class DocumentFixerTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        StandardIngestor().ingest_pdf('gost_7_32_2017')
        cls.pipeline = DocumentPipeline(llm_provider=FakeProvider(), registry=StandardRegistry())

    def test_apply_fixes_corrects_common_issues(self) -> None:
        document = DocumentInput.model_validate_json(
            Path('tests/fixtures/documents/demo_appendix_heading.json').read_text(encoding='utf-8-sig')
        )
        result = self.pipeline.analyze_document(document)

        fixed = apply_fixes(document, result.issues)

        self.assertTrue(any(section.title == '\u0412\u044b\u0432\u043e\u0434\u044b' for section in fixed.sections))
        self.assertTrue(any('\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435' in section.title for section in fixed.sections))
        self.assertTrue(any('\u0421\u041f\u0418\u0421\u041e\u041a \u0418\u0421\u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u041d\u041d\u042b\u0425 \u0418\u0421\u0422\u041e\u0427\u041d\u0418\u041a\u041e\u0412' in section.title for section in fixed.sections))
        first_section = fixed.sections[0]
        self.assertIn('\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0438', first_section.text.lower())

    def test_build_corrected_docx_returns_valid_docx_bytes(self) -> None:
        document = DocumentInput.model_validate_json(
            Path('tests/fixtures/documents/demo_figure_caption.json').read_text(encoding='utf-8-sig')
        )
        result = self.pipeline.analyze_document(document)
        fixed = apply_fixes(document, result.issues)

        content = build_corrected_docx(fixed)

        self.assertTrue(content.startswith(b'PK'))
        with zipfile.ZipFile(BytesIO(content)) as archive:
            xml = archive.read('word/document.xml').decode('utf-8')
        self.assertNotIn('\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 -', xml)
        self.assertIn('\u0421\u041f\u0418\u0421\u041e\u041a \u0418\u0421\u041f\u041e\u041b\u042c\u0417\u041e\u0412\u0410\u041d\u041d\u042b\u0425 \u0418\u0421\u0422\u041e\u0427\u041d\u0418\u041a\u041e\u0412', xml)


    def test_apply_fixes_normalizes_abbreviated_figure_caption(self) -> None:
        document = DocumentInput(
            document_id='doc_abbrev_fix',
            standard_id='gost_7_32_2017',
            meta={'filename': 'demo.docx', 'title': 'Demo', 'language': 'ru'},
            sections=[{'id': 'sec_1', 'number': '1', 'title': '\u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435', 'level': 1, 'text': ''}],
            paragraphs=[{'id': 'p_1', 'section_id': 'sec_1', 'text': '\u0421\u043c. \u0440\u0438\u0441. 1.', 'position': {'page': 1, 'paragraph_index': 1}}],
            figures=[{'id': 'fig_1', 'caption': '\u0420\u0438\u0441. 1 \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b', 'position': {'page': 1, 'paragraph_index': 2}}],
        )
        result = self.pipeline.analyze_document(document)

        fixed = apply_fixes(document, result.issues)

        self.assertEqual(fixed.figures[0].caption, '\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b')


    def test_apply_fixes_does_not_synthesize_missing_figure_caption(self) -> None:
        document = DocumentInput.model_validate_json(
            Path('tests/fixtures/documents/demo_figure_caption.json').read_text(encoding='utf-8-sig')
        )
        result = self.pipeline.analyze_document(document)

        fixed = apply_fixes(document, result.issues)

        self.assertEqual(fixed.figures[0].caption, '')


    def test_apply_fixes_to_source_docx_centers_figure_caption(self) -> None:
        source = WordDocument()
        source.add_paragraph('\u0412 \u0442\u0435\u043a\u0441\u0442\u0435 \u0434\u0430\u043d\u0430 \u0441\u0441\u044b\u043b\u043a\u0430 \u043d\u0430 \u0440\u0438\u0441\u0443\u043d\u043e\u043a 1.')
        caption = source.add_paragraph('\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b')
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='figure_alignment.docx',
            standard_id='gost_7_32_2017',
            document_id='figure_alignment',
        )
        result = self.pipeline.analyze_document(parsed)

        self.assertIn('figure_caption_not_centered', [issue.subtype for issue in result.issues])

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))
        fixed_caption = next(paragraph for paragraph in fixed_doc.paragraphs if '\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b' in paragraph.text)

        self.assertEqual(fixed_caption.alignment, WD_ALIGN_PARAGRAPH.CENTER)


    def test_apply_fixes_to_source_docx_keeps_centering_when_caption_text_changes(self) -> None:
        source = WordDocument()
        source.add_paragraph('1 \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435')
        source.add_paragraph('\u0412 \u0442\u0435\u043a\u0441\u0442\u0435 \u0434\u0430\u043d\u0430 \u0441\u0441\u044b\u043b\u043a\u0430 \u043d\u0430 \u0440\u0438\u0441\u0443\u043d\u043e\u043a 1.')
        caption = source.add_paragraph('\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b.')
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='figure_alignment_and_format.docx',
            standard_id='gost_7_32_2017',
            document_id='figure_alignment_and_format',
        )
        result = self.pipeline.analyze_document(parsed)

        self.assertIn('figure_caption_not_centered', [issue.subtype for issue in result.issues])
        self.assertIn('invalid_figure_caption_format', [issue.subtype for issue in result.issues])

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))
        fixed_caption = next(paragraph for paragraph in fixed_doc.paragraphs if '\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1' in paragraph.text)

        self.assertEqual(fixed_caption.alignment, WD_ALIGN_PARAGRAPH.CENTER)


    def test_apply_fixes_to_source_docx_updates_multiple_figure_captions(self) -> None:
        source = WordDocument()
        source.add_paragraph('1 \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435')
        source.add_paragraph('\u0412 \u0442\u0435\u043a\u0441\u0442\u0435 \u0434\u0430\u043d\u0430 \u0441\u0441\u044b\u043b\u043a\u0430 \u043d\u0430 \u0440\u0438\u0441\u0443\u043d\u043e\u043a 1.')
        first = source.add_paragraph('\u0420\u0438\u0441. 1 \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b.')
        first.alignment = WD_ALIGN_PARAGRAPH.LEFT
        source.add_paragraph('\u0412 \u0442\u0435\u043a\u0441\u0442\u0435 \u0434\u0430\u043d\u0430 \u0441\u0441\u044b\u043b\u043a\u0430 \u043d\u0430 \u0440\u0438\u0441\u0443\u043d\u043e\u043a 2.')
        second = source.add_paragraph('\u0440\u0438\u0441 2 - \u0421\u0445\u0435\u043c\u0430 \u043c\u043e\u0434\u0443\u043b\u044f.')
        second.alignment = WD_ALIGN_PARAGRAPH.LEFT

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='multiple_figures.docx',
            standard_id='gost_7_32_2017',
            document_id='multiple_figures',
        )
        result = self.pipeline.analyze_document(parsed)

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))
        captions = [paragraph for paragraph in fixed_doc.paragraphs if '\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1' in paragraph.text or '\u0420\u0438\u0441\u0443\u043d\u043e\u043a 2' in paragraph.text]

        self.assertEqual(len(captions), 2)
        self.assertIn('\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b', captions[0].text)
        self.assertIn('\u0420\u0438\u0441\u0443\u043d\u043e\u043a 2 - \u0421\u0445\u0435\u043c\u0430 \u043c\u043e\u0434\u0443\u043b\u044f', captions[1].text)
        self.assertFalse(captions[0].text.endswith('.'))
        self.assertFalse(captions[1].text.endswith('.'))
        self.assertEqual(captions[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(captions[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)


    def test_apply_fixes_to_source_docx_renumbers_section_headings(self) -> None:
        source = WordDocument()
        source.add_paragraph('2 Р’РІРµРґРµРЅРёРµ')
        source.add_paragraph('РўРµРєСЃС‚ РІРІРµРґРµРЅРёСЏ')
        source.add_paragraph('2.2 РћСЃРЅРѕРІРЅР°СЏ С‡Р°СЃС‚СЊ')
        source.add_paragraph('РўРµРєСЃС‚ РѕСЃРЅРѕРІРЅРѕР№ С‡Р°СЃС‚Рё')
        source.add_paragraph('4 Р—Р°РєР»СЋС‡РµРЅРёРµ')
        source.add_paragraph('РўРµРєСЃС‚ Р·Р°РєР»СЋС‡РµРЅРёСЏ')
        source.add_paragraph('5 РЎРїРёСЃРѕРє РёСЃРїРѕР»СЊР·РѕРІР°РЅРЅС‹С… РёСЃС‚РѕС‡РЅРёРєРѕРІ')
        source.add_paragraph('[1] РСЃС‚РѕС‡РЅРёРє')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='section_renumbering.docx',
            standard_id='gost_7_32_2017',
            document_id='section_renumbering',
        )
        result = self.pipeline.analyze_document(parsed)

        self.assertIn('numbering_error', [issue.subtype for issue in result.issues])

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))
        headings = [paragraph.text for paragraph in fixed_doc.paragraphs if paragraph.text and paragraph.text[0].isdigit()]

        self.assertIn('1 Р’РІРµРґРµРЅРёРµ', headings)
        self.assertIn('1.1 РћСЃРЅРѕРІРЅР°СЏ С‡Р°СЃС‚СЊ', headings)
        self.assertIn('2 Р—Р°РєР»СЋС‡РµРЅРёРµ', headings)


    def test_apply_fixes_to_source_docx_fixes_formula_layout_number_and_reference(self) -> None:
        source = WordDocument()
        source.add_paragraph('В тексте далее используется по формуле 1.')
        formula = source.add_paragraph('F(x) = a + b + c 1')
        formula.alignment = WD_ALIGN_PARAGRAPH.LEFT
        source.add_paragraph('где: a - коэффициент')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='formula_fix.docx',
            standard_id='gost_7_32_2017',
            document_id='formula_fix',
        )
        result = self.pipeline.analyze_document(parsed)

        subtypes = {issue.subtype for issue in result.issues}
        self.assertIn('formula_not_centered', subtypes)
        self.assertIn('formula_number_format_error', subtypes)
        self.assertIn('formula_where_colon', subtypes)

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        fixed_formula = next(paragraph for paragraph in fixed_doc.paragraphs if 'F(x)' in paragraph.text)
        fixed_reference = fixed_doc.paragraphs[0]
        fixed_explanation = fixed_doc.paragraphs[2]

        self.assertEqual(fixed_formula.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(fixed_formula.text, 'F(x) = a + b + c (1)')
        self.assertIn('формуле (1)', fixed_reference.text)
        self.assertEqual(fixed_explanation.text, 'где a - коэффициент')

    
    def test_apply_fixes_to_source_docx_moves_inline_formula_to_separate_centered_paragraph(self) -> None:
        source = WordDocument()
        source.add_paragraph('\u0420\u0430\u0441\u0447\u0435\u0442 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u044e\u0442 \u043f\u043e \u0432\u044b\u0440\u0430\u0436\u0435\u043d\u0438\u044e F(x) = a + b + c.')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='inline_formula.docx',
            standard_id='gost_7_32_2017',
            document_id='inline_formula',
        )
        result = self.pipeline.analyze_document(parsed)

        subtypes = {issue.subtype for issue in result.issues}
        self.assertIn('formula_not_standalone', subtypes)
        self.assertIn('missing_formula_number', subtypes)

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        self.assertEqual(fixed_doc.paragraphs[0].text, '\u0420\u0430\u0441\u0447\u0435\u0442 \u0432\u044b\u043f\u043e\u043b\u043d\u044f\u044e\u0442 \u043f\u043e \u0432\u044b\u0440\u0430\u0436\u0435\u043d\u0438\u044e')
        self.assertEqual(fixed_doc.paragraphs[1].text, 'F(x) = a + b + c (1)')
        self.assertEqual(fixed_doc.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_apply_fixes_to_source_docx_keeps_multiple_formulas_and_numbers_them_sequentially(self) -> None:
        source = WordDocument()
        source.add_paragraph('\u042d\u043d\u0435\u0440\u0433\u0438\u044f \u0432\u044b\u0447\u0438\u0441\u043b\u044f\u0435\u0442\u0441\u044f \u043f\u043e \u0444\u043e\u0440\u043c\u0443\u043b\u0435: E = mc2')
        source.add_paragraph('\u0433\u0434\u0435: E \u2014 \u044d\u043d\u0435\u0440\u0433\u0438\u044f; m \u2014 \u043c\u0430\u0441\u0441\u0430; c \u2014 \u0441\u043a\u043e\u0440\u043e\u0441\u0442\u044c \u0441\u0432\u0435\u0442\u0430 \u0432 \u0432\u0430\u043a\u0443\u0443\u043c\u0435.')
        source.add_paragraph('\u0421\u043b\u0435\u0434\u0443\u044e\u0449\u0435\u0435 \u0441\u043e\u043e\u0442\u043d\u043e\u0448\u0435\u043d\u0438\u0435 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u0443\u0435\u0442\u0441\u044f \u043f\u0440\u0438 \u0432\u044b\u0447\u0438\u0441\u043b\u0435\u043d\u0438\u0438 \u0441\u0440\u0435\u0434\u043d\u0435\u0439 \u0432\u0435\u043b\u0438\u0447\u0438\u043d\u044b \u043f\u043e \u0444\u043e\u0440\u043c\u0443\u043b\u0435: x = (x1 + x2 + x3) / n')
        source.add_paragraph('\u0433\u0434\u0435 n \u2014 \u043a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e \u043d\u0430\u0431\u043b\u044e\u0434\u0435\u043d\u0438\u0439, x1, x2, x3 \u2014 \u044d\u043b\u0435\u043c\u0435\u043d\u0442\u044b \u0432\u044b\u0431\u043e\u0440\u043a\u0438.')
        source.add_paragraph('D = (x1 + x2) / n')
        source.add_paragraph('Q = a / b')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='formula_sequence.docx',
            standard_id='gost_7_32_2017',
            document_id='formula_sequence',
        )
        result = self.pipeline.analyze_document(parsed)
        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        texts = [paragraph.text for paragraph in fixed_doc.paragraphs]
        self.assertIn('\u042d\u043d\u0435\u0440\u0433\u0438\u044f \u0432\u044b\u0447\u0438\u0441\u043b\u044f\u0435\u0442\u0441\u044f \u043f\u043e \u0444\u043e\u0440\u043c\u0443\u043b\u0435 (1)', texts)
        self.assertIn('E = mc2 (1)', texts)
        self.assertIn('\u0433\u0434\u0435 E \u2014 \u044d\u043d\u0435\u0440\u0433\u0438\u044f; m \u2014 \u043c\u0430\u0441\u0441\u0430; c \u2014 \u0441\u043a\u043e\u0440\u043e\u0441\u0442\u044c \u0441\u0432\u0435\u0442\u0430 \u0432 \u0432\u0430\u043a\u0443\u0443\u043c\u0435.', texts)
        self.assertIn('\u0421\u043b\u0435\u0434\u0443\u044e\u0449\u0435\u0435 \u0441\u043e\u043e\u0442\u043d\u043e\u0448\u0435\u043d\u0438\u0435 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u0443\u0435\u0442\u0441\u044f \u043f\u0440\u0438 \u0432\u044b\u0447\u0438\u0441\u043b\u0435\u043d\u0438\u0438 \u0441\u0440\u0435\u0434\u043d\u0435\u0439 \u0432\u0435\u043b\u0438\u0447\u0438\u043d\u044b \u043f\u043e \u0444\u043e\u0440\u043c\u0443\u043b\u0435 (2)', texts)
        self.assertIn('x = (x1 + x2 + x3) / n (2)', texts)
        self.assertIn('D = (x1 + x2) / n (3)', texts)
        self.assertIn('Q = a / b (4)', texts)
        self.assertEqual(sum(1 for value in texts if 'mc2' in value), 1)
        self.assertEqual(sum(1 for value in texts if 'x1 + x2 + x3' in value), 1)
        self.assertEqual(sum(1 for value in texts if 'Q = a / b' in value), 1)

    def test_apply_fixes_to_source_docx_updates_typography_and_margins(self) -> None:
        source = WordDocument()
        section = source.sections[0]
        section.page_width = Mm(180)
        section.page_height = Mm(260)
        section.left_margin = Mm(25)
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        heading = source.add_paragraph('1 РћСЃРЅРѕРІРЅР°СЏ С‡Р°СЃС‚СЊ')
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(14)
        heading_run.font.name = 'Calibri'
        heading_run.font.color.rgb = RGBColor(0, 0, 255)
        body = source.add_paragraph('??? ???????? ????? ? ????????? ???????????? ??????????? ??? ????????.')
        body.paragraph_format.first_line_indent = Mm(0)
        body.paragraph_format.line_spacing = 1.0
        run = body.runs[0]
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        second_body = source.add_paragraph('?????? ?????? ????? ??? ??????? ??? ??????? ????? ??????.')
        second_body.paragraph_format.first_line_indent = Mm(0)
        second_body.paragraph_format.line_spacing = 1.0
        second_run = second_body.runs[0]
        second_run.font.size = Pt(10)
        second_run.font.name = 'Calibri'
        second_run.font.bold = True
        second_run.font.color.rgb = RGBColor(128, 128, 128)

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='typography_fix.docx',
            standard_id='gost_7_32_2017',
            document_id='typography_fix',
        )
        result = self.pipeline.analyze_document(parsed)

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))
        fixed_section = fixed_doc.sections[0]
        fixed_heading = next(paragraph for paragraph in fixed_doc.paragraphs if '1 РћСЃРЅРѕРІРЅР°СЏ С‡Р°СЃС‚СЊ' in paragraph.text)
        fixed_heading_run = fixed_heading.runs[0]
        fixed_body = next(paragraph for paragraph in fixed_doc.paragraphs if '??? ???????? ?????' in paragraph.text)
        fixed_run = fixed_body.runs[0]
        fixed_second_body = next(paragraph for paragraph in fixed_doc.paragraphs if '?????? ??????' in paragraph.text)
        fixed_second_run = fixed_second_body.runs[0]

        self.assertAlmostEqual(fixed_section.page_width.mm, 210.0, places=1)
        self.assertAlmostEqual(fixed_section.page_height.mm, 297.0, places=1)
        self.assertAlmostEqual(fixed_section.left_margin.mm, 30.0, places=1)
        self.assertAlmostEqual(fixed_section.right_margin.mm, 15.0, places=1)
        self.assertAlmostEqual(fixed_section.top_margin.mm, 20.0, places=1)
        self.assertAlmostEqual(fixed_section.bottom_margin.mm, 20.0, places=1)
        self.assertAlmostEqual(fixed_heading_run.font.size.pt, 12.0, places=1)
        self.assertEqual(fixed_heading_run.font.name, 'Times New Roman')
        self.assertEqual(str(fixed_heading_run.font.color.rgb), '000000')
        self.assertTrue(bool(fixed_heading_run.font.bold))
        self.assertAlmostEqual(fixed_body.paragraph_format.first_line_indent.mm, 12.5, places=1)
        self.assertAlmostEqual(float(fixed_body.paragraph_format.line_spacing), 1.5, places=1)
        self.assertAlmostEqual(fixed_run.font.size.pt, 12.0, places=1)
        self.assertEqual(fixed_run.font.name, 'Times New Roman')
        self.assertFalse(bool(fixed_run.font.bold))
        self.assertEqual(str(fixed_run.font.color.rgb), '000000')
        self.assertAlmostEqual(fixed_second_body.paragraph_format.first_line_indent.mm, 12.5, places=1)
        self.assertAlmostEqual(float(fixed_second_body.paragraph_format.line_spacing), 1.5, places=1)
        self.assertAlmostEqual(fixed_second_run.font.size.pt, 12.0, places=1)
        self.assertEqual(fixed_second_run.font.name, 'Times New Roman')
        self.assertFalse(bool(fixed_second_run.font.bold))
        self.assertEqual(str(fixed_second_run.font.color.rgb), '000000')


    def test_apply_fixes_to_source_docx_moves_table_caption_and_cleans_headers_in_appendix(self) -> None:
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P

        source = WordDocument()
        source.add_paragraph('\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435 \u0410 - \u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044b')
        source.add_paragraph('\u0412 \u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0438 \u0410 \u043f\u0440\u0438\u0432\u0435\u0434\u0435\u043d\u0430 \u0442\u0430\u0431\u043b\u0438\u0446\u0430 \u0410.1.')
        table = source.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = '\u041f\u043e\u043a\u0430\u0437\u0430\u0442\u0435\u043b\u044c.'
        table.rows[0].cells[1].text = '\u0417\u043d\u0430\u0447\u0435\u043d\u0438\u0435.'
        table.rows[1].cells[0].text = '\u0421\u043a\u043e\u0440\u043e\u0441\u0442\u044c'
        table.rows[1].cells[1].text = '10'
        source.add_paragraph('\u0422\u0430\u0431\u043b\u0438\u0446\u0430 1 - \u0414\u0430\u043d\u043d\u044b\u0435.')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='appendix_table.docx',
            standard_id='gost_7_32_2017',
            document_id='appendix_table',
        )
        result = self.pipeline.analyze_document(parsed)

        self.assertIn('table_caption_below_table', [issue.subtype for issue in result.issues])
        self.assertIn('invalid_table_header_punctuation', [issue.subtype for issue in result.issues])
        self.assertIn('appendix_table_caption_format', [issue.subtype for issue in result.issues])

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        body_kinds = []
        body_texts = []
        for child in fixed_doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = next((p for p in fixed_doc.paragraphs if p._p is child), None)
                body_kinds.append('p')
                body_texts.append(paragraph.text if paragraph is not None else '')
            elif isinstance(child, CT_Tbl):
                body_kinds.append('tbl')
                body_texts.append('<table>')

        table_pos = body_kinds.index('tbl')
        caption_pos = body_texts.index('\u0422\u0430\u0431\u043b\u0438\u0446\u0430 \u0410.1 - \u0414\u0430\u043d\u043d\u044b\u0435')
        self.assertLess(caption_pos, table_pos)
        self.assertEqual(fixed_doc.tables[0].rows[0].cells[0].text, '\u041f\u043e\u043a\u0430\u0437\u0430\u0442\u0435\u043b\u044c')
        self.assertEqual(fixed_doc.tables[0].rows[0].cells[1].text, '\u0417\u043d\u0430\u0447\u0435\u043d\u0438\u0435')


    def test_apply_fixes_to_source_docx_formats_headings_appendix_and_table_header_alignment(self) -> None:
        source = WordDocument()
        main_heading = source.add_paragraph('1 \u041e\u0441\u043d\u043e\u0432\u043d\u0430\u044f \u0447\u0430\u0441\u0442\u044c')
        main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        main_heading.paragraph_format.first_line_indent = Mm(0)
        main_heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)

        source.add_paragraph('\u0422\u0435\u043a\u0441\u0442 \u043e\u0441\u043d\u043e\u0432\u043d\u043e\u0439 \u0447\u0430\u0441\u0442\u0438')

        refs_heading = source.add_paragraph('3 \u0421\u043f\u0438\u0441\u043e\u043a \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u0438\u0441\u0442\u043e\u0447\u043d\u0438\u043a\u043e\u0432')
        refs_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        refs_heading.paragraph_format.first_line_indent = Mm(12.5)
        refs_heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)

        appendix_heading = source.add_paragraph('\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435 \u0410 - \u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044b')
        appendix_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = source.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = '\u041f\u043e\u043a\u0430\u0437\u0430\u0442\u0435\u043b\u044c'
        table.rows[0].cells[1].text = '\u0417\u043d\u0430\u0447\u0435\u043d\u0438\u0435'
        table.rows[1].cells[0].text = '\u0421\u043a\u043e\u0440\u043e\u0441\u0442\u044c'
        table.rows[1].cells[1].text = '10'
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        source.add_paragraph('\u0422\u0430\u0431\u043b\u0438\u0446\u0430 1 - \u0414\u0430\u043d\u043d\u044b\u0435')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='heading_table_layout.docx',
            standard_id='gost_7_32_2017',
            document_id='heading_table_layout',
        )
        result = self.pipeline.analyze_document(parsed)

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, result.issues)
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        fixed_main = next(paragraph for paragraph in fixed_doc.paragraphs if paragraph.text == '1 \u041e\u0441\u043d\u043e\u0432\u043d\u0430\u044f \u0447\u0430\u0441\u0442\u044c')
        fixed_refs = next(paragraph for paragraph in fixed_doc.paragraphs if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and paragraph.text and paragraph.text == paragraph.text.upper() and '\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415' not in paragraph.text)
        fixed_appendix = next(paragraph for paragraph in fixed_doc.paragraphs if paragraph.text == '\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415 \u0410')
        fixed_appendix_title = next(paragraph for paragraph in fixed_doc.paragraphs if paragraph.text == '\u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044b')

        self.assertEqual(fixed_main.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertAlmostEqual(fixed_main.paragraph_format.first_line_indent.mm, 12.5, places=1)
        self.assertEqual(str(fixed_main.runs[0].font.color.rgb), '000000')
        self.assertTrue(bool(fixed_main.runs[0].font.bold))

        self.assertEqual(fixed_refs.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertAlmostEqual(fixed_refs.paragraph_format.first_line_indent.mm, 0.0, places=1)
        self.assertEqual(str(fixed_refs.runs[0].font.color.rgb), '000000')
        self.assertTrue(bool(fixed_refs.runs[0].font.bold))

        self.assertEqual(fixed_appendix.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(fixed_appendix_title.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(bool(fixed_appendix.runs[0].font.bold))
        self.assertTrue(bool(fixed_appendix_title.runs[0].font.bold))

        for cell in fixed_doc.tables[0].rows[0].cells:
            for paragraph in cell.paragraphs:
                self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_apply_fixes_to_source_docx_does_not_rewrite_native_math_formula_text(self) -> None:
        source = WordDocument()
        source.add_paragraph('?????? ????????? ?? ??????? (4).')
        formula = source.add_paragraph(r'D = \frac{\sum_{i=1}^{n} (x_i - \bar{x})^2}{n} (4)')
        source.add_paragraph('??????????? ???????')

        buffer = BytesIO()
        source.save(buffer)
        source_bytes = buffer.getvalue()

        parsed = parse_docx_to_document(
            source_bytes,
            filename='native_math_guard.docx',
            standard_id='gost_7_32_2017',
            document_id='native_math_guard',
        )
        parsed.meta.extras['docx_formulas'] = [
            {
                'paragraph_index': 2,
                'text': r'D = \frac{\sum_{i=1}^{n} (x_i - \bar{x})^2}{n} (4)',
                'alignment': 'left',
                'has_math_xml': True,
                'is_standalone': True,
                'equation_number': '4',
                'raw_equation_number': '',
                'prev_blank': True,
                'next_blank': True,
                'prev_text': '?????? ????????? ?? ??????? (4).',
                'next_text': '??????????? ???????',
                'next_paragraph_index': 3,
                'section_id': None,
                'section_title': '',
                'appendix_letter': '',
            }
        ]

        fixed_bytes = apply_fixes_to_source_docx(source_bytes, parsed, [])
        fixed_doc = WordDocument(BytesIO(fixed_bytes))

        self.assertEqual(fixed_doc.paragraphs[1].text, r'D = \frac{\sum_{i=1}^{n} (x_i - \bar{x})^2}{n} (4)')

if __name__ == '__main__':
    unittest.main()


