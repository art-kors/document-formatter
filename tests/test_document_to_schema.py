from io import BytesIO
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm

from app.parsing.document_to_schema import parse_docx_to_document, parse_text_to_document


def _add_page_field(paragraph) -> None:
    begin_run = paragraph.add_run()._r
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    begin_run.append(fld_begin)

    instr_run = paragraph.add_run()._r
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    instr_run.append(instr)

    separate_run = paragraph.add_run()._r
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    separate_run.append(fld_separate)

    text_run = paragraph.add_run()._r
    page_text = OxmlElement('w:t')
    page_text.text = '1'
    text_run.append(page_text)

    end_run = paragraph.add_run()._r
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run.append(fld_end)


class DocumentToSchemaTests(unittest.TestCase):
    def test_parse_text_to_document_extracts_sections_paragraphs_and_captions(self) -> None:
        text = (
            "1 \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435\n\n"
            "\u0412\u043e \u0432\u0432\u0435\u0434\u0435\u043d\u0438\u0438 \u043e\u043f\u0438\u0441\u044b\u0432\u0430\u0435\u0442\u0441\u044f \u0446\u0435\u043b\u044c \u0440\u0430\u0431\u043e\u0442\u044b.\n\n"
            "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b\n\n"
            "\u0422\u0430\u0431\u043b\u0438\u0446\u0430 1 - \u0421\u0440\u0430\u0432\u043d\u0435\u043d\u0438\u0435 \u043c\u0435\u0442\u043e\u0434\u043e\u0432\n\n"
            "6 \u0421\u043f\u0438\u0441\u043e\u043a \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u0438\u0441\u0442\u043e\u0447\u043d\u0438\u043a\u043e\u0432\n\n"
            "[1] \u0413\u041e\u0421\u0422 7.32-2017"
        )

        document = parse_text_to_document(
            text,
            filename="report.docx",
            standard_id="gost_7_32_2017",
            document_id="doc_test",
        )

        self.assertEqual(document.document_id, "doc_test")
        self.assertEqual(document.standard_id, "gost_7_32_2017")
        self.assertEqual(len(document.sections), 2)
        self.assertEqual(document.sections[0].number, "1")
        self.assertEqual(document.sections[1].title, "\u0421\u043f\u0438\u0441\u043e\u043a \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u0438\u0441\u0442\u043e\u0447\u043d\u0438\u043a\u043e\u0432")
        self.assertEqual(len(document.figures), 1)
        self.assertEqual(len(document.tables), 1)
        self.assertGreaterEqual(len(document.paragraphs), 2)
        self.assertEqual(document.figures[0].caption, "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 1 - \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b")
        self.assertEqual(document.tables[0].caption, "\u0422\u0430\u0431\u043b\u0438\u0446\u0430 1 - \u0421\u0440\u0430\u0432\u043d\u0435\u043d\u0438\u0435 \u043c\u0435\u0442\u043e\u0434\u043e\u0432")

    def test_parse_text_to_document_extracts_abbreviated_figure_caption(self) -> None:
        text = (
            "1 \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435\n\n"
            "\u0420\u0438\u0441. 1 \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b\n\n"
            "\u0420\u0438\u0441 2 - \u0421\u0445\u0435\u043c\u0430 \u043c\u043e\u0434\u0443\u043b\u044f"
        )

        document = parse_text_to_document(
            text,
            filename="report.txt",
            standard_id="gost_7_32_2017",
            document_id="doc_abbrev_figures",
        )

        self.assertEqual(len(document.figures), 2)
        self.assertEqual(document.figures[0].caption, "\u0420\u0438\u0441. 1 \u0410\u0440\u0445\u0438\u0442\u0435\u043a\u0442\u0443\u0440\u0430 \u0441\u0438\u0441\u0442\u0435\u043c\u044b")
        self.assertEqual(document.figures[1].caption, "\u0420\u0438\u0441 2 - \u0421\u0445\u0435\u043c\u0430 \u043c\u043e\u0434\u0443\u043b\u044f")



    def test_parse_text_to_document_extracts_caption_with_missing_last_letter(self) -> None:
        text = (
            "1 Введение\n\n"
            "Рисуно 1 - Архитектура системы\n\n"
            "Таблиц 2 - Сравнение методов"
        )

        document = parse_text_to_document(
            text,
            filename="report.txt",
            standard_id="gost_7_32_2017",
            document_id="doc_truncated_caption_prefix",
        )

        self.assertEqual(len(document.figures), 1)
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(document.figures[0].caption, "Рисуно 1 - Архитектура системы")
        self.assertEqual(document.tables[0].caption, "Таблиц 2 - Сравнение методов")


    def test_parse_text_to_document_does_not_create_fake_untitled_section(self) -> None:
        text = (
            "\u0422\u0435\u043a\u0441\u0442 \u0442\u0438\u0442\u0443\u043b\u044c\u043d\u043e\u0433\u043e \u0431\u043b\u043e\u043a\u0430\n\n"
            "2 \u041e\u0441\u043d\u043e\u0432\u043d\u0430\u044f \u0447\u0430\u0441\u0442\u044c\n\n"
            "\u0422\u0435\u043a\u0441\u0442 \u0440\u0430\u0437\u0434\u0435\u043b\u0430"
        )

        document = parse_text_to_document(
            text,
            filename="report.txt",
            standard_id="gost_7_32_2017",
            document_id="doc_without_fake_section",
        )

        self.assertEqual(len(document.sections), 1)
        self.assertEqual(document.sections[0].number, "2")
        self.assertIsNone(document.paragraphs[0].section_id)
        self.assertNotIn("\u0411\u0435\u0437 \u043d\u0430\u0437\u0432\u0430\u043d\u0438\u044f", [section.title for section in document.sections])

    def test_parse_docx_to_document_collects_alignment_metadata(self) -> None:
        source = Document()
        title = source.add_paragraph("\u041e\u0442\u0447\u0435\u0442")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        performer = source.add_paragraph("\u0412\u044b\u043f\u043e\u043b\u043d\u0438\u043b: \u0441\u0442\u0443\u0434\u0435\u043d\u0442")
        performer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        source.add_paragraph("1 \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435")
        source.add_paragraph("\u041e\u0441\u043d\u043e\u0432\u043d\u043e\u0439 \u0442\u0435\u043a\u0441\u0442 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430")

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename="report.docx",
            standard_id="gost_7_32_2017",
            document_id="doc_docx_meta",
        )

        self.assertEqual(document.meta.extras["source_format"], "docx")
        self.assertGreaterEqual(len(document.meta.extras["docx_paragraphs"]), 3)
        self.assertEqual(document.meta.extras["docx_paragraphs"][0]["alignment"], "center")
        self.assertEqual(document.meta.extras["docx_paragraphs"][1]["alignment"], "right")

    def test_parse_docx_to_document_does_not_turn_regular_paragraphs_into_sections(self) -> None:
        source = Document()
        source.add_paragraph("1 ????????")
        source.add_paragraph("??? ??????? ????? ??????, ??????? ?? ?????? ??????????? ????????? ????????.")
        source.add_paragraph("??? ???? ??????? ????? ?????? ??? ????????????? ?????.")

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename="report.docx",
            standard_id="gost_7_32_2017",
            document_id="doc_docx_no_fake_sections",
        )

        self.assertEqual(len(document.sections), 1)
        self.assertEqual(document.sections[0].number, "1")
        self.assertEqual(len(document.paragraphs), 2)

    def test_parse_docx_to_document_extracts_page_size_metadata(self) -> None:
        source = Document()
        source.sections[0].page_width = Mm(148)
        source.sections[0].page_height = Mm(210)
        source.add_paragraph("1 ????????")

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename="report.docx",
            standard_id="gost_7_32_2017",
            document_id="doc_docx_page_size",
        )

        self.assertEqual(document.meta.extras['source_format'], 'docx')
        self.assertGreaterEqual(len(document.meta.extras['docx_sections']), 1)
        self.assertAlmostEqual(document.meta.extras['docx_sections'][0]['page_width_mm'], 148.0, places=1)
        self.assertAlmostEqual(document.meta.extras['docx_sections'][0]['page_height_mm'], 210.0, places=1)
        self.assertIsNotNone(document.meta.extras['docx_sections'][0]['left_margin_mm'])



    def test_parse_docx_to_document_extracts_page_numbering_metadata(self) -> None:
        source = Document()
        source.sections[0].different_first_page_header_footer = True
        footer_paragraph = source.sections[0].footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_field(footer_paragraph)
        header_paragraph = source.sections[0].header.paragraphs[0]
        _add_page_field(header_paragraph)
        source.add_paragraph("1 ????????")

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename="report.docx",
            standard_id="gost_7_32_2017",
            document_id="doc_docx_page_numbering",
        )

        numbering = document.meta.extras['docx_page_numbering']
        self.assertEqual(len(numbering), 1)
        self.assertTrue(numbering[0]['default_footer_has_page_field'])
        self.assertEqual(numbering[0]['default_footer_alignment'], 'center')
        self.assertTrue(numbering[0]['different_first_page'])
        self.assertFalse(numbering[0]['first_footer_has_page_field'])
        self.assertTrue(numbering[0]['default_header_has_page_field'])


    def test_parse_docx_to_document_does_not_treat_plain_text_as_formula(self) -> None:
        source = Document()
        source.add_paragraph('3. ?????????, ???????? ?? ?????????? ???.')
        source.add_paragraph('? ???? ??????? ??????????????? ????????? ?????????? ? ??????? ??????????.')

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename='report.docx',
            standard_id='gost_7_32_2017',
            document_id='doc_plain_text_not_formula',
        )

        self.assertEqual(document.meta.extras.get('docx_formulas', []), [])


    def test_parse_docx_to_document_does_not_treat_parameter_sentence_as_formula(self) -> None:
        source = Document()
        source.add_paragraph(
            '\u0433\u0434\u0435 Imax \u2014 \u043f\u0438\u043a\u043e\u0432\u044b\u0439 \u0442\u043e\u043a, '
            'k \u2014 \u043a\u043e\u044d\u0444\u0444\u0438\u0446\u0438\u0435\u043d\u0442, '
            'P \u2014 \u043f\u0440\u0435\u0434\u0435\u043b\u044c\u043d\u0430\u044f \u0441\u0440\u0435\u0434\u043d\u044f\u044f \u043c\u043e\u0449\u043d\u043e\u0441\u0442\u044c, '
            'f \u2014 \u0447\u0430\u0441\u0442\u043e\u0442\u0430 \u0441\u043b\u0435\u0434\u043e\u0432\u0430\u043d\u0438\u044f \u0438\u043c\u043f\u0443\u043b\u044c\u0441\u043e\u0432. '
            'k=0.61 \u043f\u0440\u0438 P=5500 \u0412\u0442, k=0.52 \u043f\u0440\u0438 P>5500 \u0412\u0442.'
        )

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename='report.docx',
            standard_id='gost_7_32_2017',
            document_id='doc_parameter_sentence_not_formula',
        )

        self.assertEqual(document.meta.extras.get('docx_formulas', []), [])

    def test_parse_docx_to_document_extracts_formula_metadata(self) -> None:
        source = Document()
        source.add_paragraph('1 ????????')
        source.add_paragraph('? ??????? (1) ???????? ??????????? ??????????.')
        formula = source.add_paragraph('E = mc2 (1)')
        formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.add_paragraph('??? E ? ???????')

        buffer = BytesIO()
        source.save(buffer)

        document = parse_docx_to_document(
            buffer.getvalue(),
            filename='report.docx',
            standard_id='gost_7_32_2017',
            document_id='doc_formula_meta',
        )

        formulas = document.meta.extras.get('docx_formulas', [])
        self.assertEqual(len(formulas), 1)
        self.assertEqual(formulas[0]['equation_number'], '1')
        self.assertEqual(formulas[0]['alignment'], 'center')
        self.assertEqual(formulas[0]['next_text'], '??? E ? ???????')


if __name__ == "__main__":
    unittest.main()
