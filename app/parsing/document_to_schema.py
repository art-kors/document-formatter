import re
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from app.schemas.document import DocumentInput, DocumentMeta, FigureItem, Paragraph, Position, Section, TableItem


SECTION_PATTERN = re.compile(r'^(?P<number>\d+(?:\.\d+)*)\s+(?P<title>[^\n]{1,200})$')
APPENDIX_PATTERN = re.compile('^(?:\u041f\u0420\u0418\u041b\u041e\u0416\u0415\u041d\u0418\u0415|\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435)\\s+(?P<number>[A-Z\u0410-\u042f\u0401])(?:\\s*[\u2013-]?\\s*(?P<title>.*))?$')
FIGURE_PATTERN = re.compile(
    '^(?:\u0420\u0438\u0441\u0443\u043d\u043e\u043a|\u0420\u0438\u0441\\.?|\u0440\u0438\u0441\\.?)\\s*(?P<number>\\d+(?:\\.\\d+)*)\\s*[\u2014\u2013-]?\\s*(?P<title>.*)$',
    re.IGNORECASE,
)
TABLE_PATTERN = re.compile('^\u0422\u0430\u0431\u043b\u0438\u0446\u0430\\s+(?P<number>\\d+)\\s*[\u2013-]?\\s*(?P<title>.*)$', re.IGNORECASE)
_HEADING_STYLE_HINTS = ('heading', '\u0437\u0430\u0433\u043e\u043b\u043e\u0432\u043e\u043a')

_CANONICAL_UNNUMBERED_HEADINGS = {
    '\u0432\u0432\u0435\u0434\u0435\u043d\u0438\u0435',
    '\u0437\u0430\u043a\u043b\u044e\u0447\u0435\u043d\u0438\u0435',
    '\u0441\u043e\u0434\u0435\u0440\u0436\u0430\u043d\u0438\u0435',
    '\u043e\u0433\u043b\u0430\u0432\u043b\u0435\u043d\u0438\u0435',
    '\u0440\u0435\u0444\u0435\u0440\u0430\u0442',
    '\u0441\u043f\u0438\u0441\u043e\u043a \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u0438\u0441\u0442\u043e\u0447\u043d\u0438\u043a\u043e\u0432',
    '\u0441\u043f\u0438\u0441\u043e\u043a \u043b\u0438\u0442\u0435\u0440\u0430\u0442\u0443\u0440\u044b',
    '\u0431\u0438\u0431\u043b\u0438\u043e\u0433\u0440\u0430\u0444\u0438\u044f',
}


def parse_text_to_document(
    text: str,
    *,
    filename: str,
    standard_id: str,
    document_id: Optional[str] = None,
) -> DocumentInput:
    blocks = _split_blocks(text)
    sections: List[Section] = []
    paragraphs: List[Paragraph] = []
    figures: List[FigureItem] = []
    tables: List[TableItem] = []

    current_section: Optional[Section] = None

    for paragraph_index, block in enumerate(blocks, start=1):
        heading = _parse_heading(block)
        if heading is not None:
            if current_section is not None:
                current_section.text = current_section.text.strip()
            current_section = Section(**heading)
            sections.append(current_section)
            continue

        if current_section is not None:
            current_section.text = (current_section.text + "\n\n" + block).strip() if current_section.text else block

        figure = _parse_figure(block, len(figures) + 1, paragraph_index)
        if figure is not None:
            figures.append(figure)
            continue

        table = _parse_table(block, len(tables) + 1, paragraph_index)
        if table is not None:
            tables.append(table)
            continue

        paragraphs.append(
            Paragraph(
                id=f"p_{len(paragraphs) + 1}",
                section_id=current_section.id if current_section is not None else None,
                text=block,
                position=Position(page=None, paragraph_index=paragraph_index),
            )
        )

    if current_section is not None:
        current_section.text = current_section.text.strip()

    title = next((section.title for section in sections if section.title), Path(filename).stem)
    return DocumentInput(
        document_id=document_id or _build_document_id(filename),
        standard_id=standard_id,
        meta=DocumentMeta(filename=filename, title=title, language='ru', extras={'source_format': 'text'}),
        sections=sections,
        paragraphs=paragraphs,
        tables=tables,
        figures=figures,
    )

def parse_docx_to_document(
    file_bytes: bytes,
    *,
    filename: str,
    standard_id: str,
    document_id: Optional[str] = None,
) -> DocumentInput:
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    sections: List[Section] = []
    paragraphs: List[Paragraph] = []
    figures: List[FigureItem] = []
    tables: List[TableItem] = []
    paragraph_meta: List[Dict[str, Any]] = []
    section_heading_meta: Dict[str, Dict[str, Any]] = {}
    current_section: Optional[Section] = None

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = ' '.join(paragraph.text.split())
        if not text:
            continue

        alignment = _alignment_name(paragraph.alignment)
        style_name = (getattr(paragraph.style, 'name', '') or '').strip()
        paragraph_meta.append(_build_paragraph_meta_entry(paragraph, index, text, alignment, style_name))

        heading = _parse_heading(text)
        if heading is None and _looks_like_word_heading(style_name, text):
            heading = _fallback_heading_from_style(text, len(sections) + 1)
        if heading is not None:
            if current_section is not None:
                current_section.text = current_section.text.strip()
            current_section = Section(**heading)
            sections.append(current_section)
            section_heading_meta[current_section.id] = {
                'paragraph_index': index,
                'text': text,
                'alignment': alignment,
                'style_name': style_name,
            }
            continue

        if current_section is not None:
            current_section.text = (current_section.text + "\n\n" + text).strip() if current_section.text else text

        figure = _parse_figure(text, len(figures) + 1, index)
        if figure is not None:
            figures.append(figure)
            continue

        table = _parse_table(text, len(tables) + 1, index)
        if table is not None:
            tables.append(table)
            continue

        paragraphs.append(
            Paragraph(
                id=f"p_{len(paragraphs) + 1}",
                section_id=current_section.id if current_section is not None else None,
                text=text,
                position=Position(page=None, paragraph_index=index),
            )
        )

    if current_section is not None:
        current_section.text = current_section.text.strip()

    docx_tables_meta = _extract_docx_tables_meta(doc, sections)
    title = next((section.title for section in sections if section.title), Path(filename).stem)
    extras = {
        'source_format': 'docx',
        'docx_paragraphs': paragraph_meta,
        'docx_table_paragraphs': _extract_table_paragraph_meta(doc),
        'docx_tables_meta': docx_tables_meta,
        'docx_formulas': _extract_docx_formula_meta(doc),
        'docx_sections': _extract_section_page_meta(doc),
        'section_headings': section_heading_meta,
        'has_tables': bool(doc.tables),
        'inline_shapes_count': len(getattr(doc, 'inline_shapes', [])),
    }
    return DocumentInput(
        document_id=document_id or _build_document_id(filename),
        standard_id=standard_id,
        meta=DocumentMeta(filename=filename, title=title, language='ru', extras=extras),
        sections=sections,
        paragraphs=paragraphs,
        tables=tables,
        figures=figures,
    )


def _split_blocks(text: str) -> List[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    return [block.strip() for block in re.split(r"\n\s*\n", normalized) if block.strip()]


def _parse_heading(block: str) -> Optional[dict]:
    single_line = ' '.join(line.strip() for line in block.splitlines() if line.strip())
    appendix_match = APPENDIX_PATTERN.match(single_line)
    if appendix_match:
        suffix = appendix_match.group('title') or ''
        title = single_line if suffix else f"\u041f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u0435 {appendix_match.group('number')}"
        return {
            'id': f"sec_appendix_{appendix_match.group('number').lower()}",
            'number': appendix_match.group('number'),
            'title': title,
            'level': 1,
            'text': '',
        }

    match = SECTION_PATTERN.match(single_line)
    if not match:
        return None

    number = match.group('number')
    title = match.group('title').strip()
    level = number.count('.') + 1
    return {
        'id': f"sec_{number.replace('.', '_')}",
        'number': number,
        'title': title,
        'level': level,
        'text': '',
    }


def _parse_figure(block: str, index: int, paragraph_index: int) -> Optional[FigureItem]:
    single_line = ' '.join(line.strip() for line in block.splitlines() if line.strip())
    match = FIGURE_PATTERN.match(single_line)
    if not match:
        return None
    return FigureItem(
        id=f'fig_{index}',
        caption=single_line,
        position=Position(page=None, paragraph_index=paragraph_index),
    )


def _parse_table(block: str, index: int, paragraph_index: int) -> Optional[TableItem]:
    single_line = ' '.join(line.strip() for line in block.splitlines() if line.strip())
    match = TABLE_PATTERN.match(single_line)
    if not match:
        return None
    return TableItem(
        id=f'tbl_{index}',
        caption=single_line,
        position=Position(page=None, paragraph_index=paragraph_index),
    )


def _looks_like_word_heading(style_name: str, text: str) -> bool:
    normalized_style = style_name.strip().lower()
    if any(hint in normalized_style for hint in _HEADING_STYLE_HINTS):
        return True

    normalized_text = ' '.join(text.lower().replace('?', '?').split())
    if normalized_text in _CANONICAL_UNNUMBERED_HEADINGS:
        return True

    return False


def _fallback_heading_from_style(text: str, sequence: int) -> dict:
    return {
        'id': f'sec_auto_{sequence}',
        'number': '',
        'title': text.strip(),
        'level': 1,
        'text': '',
    }


def _extract_section_page_meta(doc: Any) -> List[Dict[str, Any]]:
    section_meta: List[Dict[str, Any]] = []
    for index, section in enumerate(getattr(doc, 'sections', []), start=1):
        page_width_mm = _emu_to_mm(getattr(section, 'page_width', None))
        page_height_mm = _emu_to_mm(getattr(section, 'page_height', None))
        left_margin_mm = _emu_to_mm(getattr(section, 'left_margin', None))
        right_margin_mm = _emu_to_mm(getattr(section, 'right_margin', None))
        top_margin_mm = _emu_to_mm(getattr(section, 'top_margin', None))
        bottom_margin_mm = _emu_to_mm(getattr(section, 'bottom_margin', None))
        section_meta.append(
            {
                'section_index': index,
                'page_width_mm': page_width_mm,
                'page_height_mm': page_height_mm,
                'left_margin_mm': left_margin_mm,
                'right_margin_mm': right_margin_mm,
                'top_margin_mm': top_margin_mm,
                'bottom_margin_mm': bottom_margin_mm,
                'orientation': 'landscape' if page_width_mm and page_height_mm and page_width_mm > page_height_mm else 'portrait',
            }
        )
    return section_meta


def _emu_to_mm(length: Any) -> Optional[float]:
    if length is None:
        return None
    try:
        return round(float(length) / 36000.0, 1)
    except (TypeError, ValueError):
        return None


def _extract_docx_tables_meta(doc: Any, sections: List[Section]) -> List[Dict[str, Any]]:
    tables_meta: List[Dict[str, Any]] = []
    blocks = list(_iter_docx_body_blocks(doc))

    for position, block in enumerate(blocks):
        if block['type'] != 'table':
            continue
        previous_paragraph = _nearest_table_caption_block(blocks, position, step=-1)
        next_paragraph = _nearest_table_caption_block(blocks, position, step=1)
        caption_block = previous_paragraph or next_paragraph
        caption_position = 'above' if previous_paragraph else 'below' if next_paragraph else 'missing'
        section_title = block.get('section_title') or ''
        tables_meta.append({
            'table_index': block['table_index'],
            'caption_paragraph_index': caption_block.get('paragraph_index') if caption_block else None,
            'caption_position': caption_position,
            'header_cells': block.get('header_cells', []),
            'section_id': block.get('section_id'),
            'section_title': section_title,
            'appendix_letter': _extract_appendix_letter(section_title),
        })
    return tables_meta


def _nearest_table_caption_block(blocks: List[Dict[str, Any]], position: int, *, step: int) -> Optional[Dict[str, Any]]:
    index = position + step
    while 0 <= index < len(blocks):
        block = blocks[index]
        if block['type'] == 'table':
            return None
        if block['type'] == 'paragraph' and TABLE_PATTERN.match(block['text']):
            return block
        if block['type'] == 'paragraph' and _parse_heading(block['text']) is not None:
            return None
        index += step
    return None


def _extract_appendix_letter(title: str) -> Optional[str]:
    match = APPENDIX_PATTERN.match(' '.join((title or '').split()))
    return match.group('number') if match else None


def _iter_docx_body_blocks(doc: Any):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = doc.element.body
    paragraph_index = 0
    table_index = 0
    current_section_id = None
    current_section_title = None
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = ' '.join(paragraph.text.split())
            if not text:
                continue
            paragraph_index += 1
            style_name = (getattr(paragraph.style, 'name', '') or '').strip()
            heading = _parse_heading(text)
            if heading is None and _looks_like_word_heading(style_name, text):
                heading = _fallback_heading_from_style(text, 1)
            if heading is not None:
                current_section_id = heading['id']
                current_section_title = heading['title']
            yield {
                'type': 'paragraph',
                'paragraph_index': paragraph_index,
                'text': text,
                'style_name': style_name,
                'section_id': current_section_id,
                'section_title': current_section_title,
            }
        elif isinstance(child, CT_Tbl):
            table_index += 1
            table = Table(child, doc)
            header_cells = []
            if table.rows:
                for cell in table.rows[0].cells:
                    cell_text = ' '.join(par.text.strip() for par in cell.paragraphs if par.text and par.text.strip())
                    header_cells.append(cell_text)
            yield {
                'type': 'table',
                'table_index': table_index,
                'header_cells': header_cells,
                'section_id': current_section_id,
                'section_title': current_section_title,
            }

def _extract_table_paragraph_meta(doc: Any) -> List[Dict[str, Any]]:
    table_paragraphs: List[Dict[str, Any]] = []
    paragraph_index = 10000
    for table_index, table in enumerate(getattr(doc, 'tables', []), start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    text = ' '.join(paragraph.text.split())
                    if not text:
                        continue
                    table_paragraphs.append(
                        {
                            'paragraph_index': paragraph_index,
                            'text': text,
                            'alignment': _alignment_name(paragraph.alignment),
                            'style': (getattr(paragraph.style, 'name', '') or '').strip(),
                            'table_index': table_index,
                            'row_index': row_index,
                            'cell_index': cell_index,
                        }
                    )
                    paragraph_index += 1
    return table_paragraphs


def _extract_docx_formula_meta(doc: Any) -> List[Dict[str, Any]]:
    formulas: List[Dict[str, Any]] = []
    paragraphs = list(getattr(doc, 'paragraphs', []))
    current_section_title = ''
    current_section_id = None

    for index, paragraph in enumerate(paragraphs, start=1):
        text = ' '.join(paragraph.text.split())
        if _starts_with_formula_explanation(text):
            continue
        style_name = (getattr(paragraph.style, 'name', '') or '').strip()
        heading = _parse_heading(text)
        if heading is None and _looks_like_word_heading(style_name, text):
            heading = _fallback_heading_from_style(text, 1)
        if heading is not None:
            current_section_id = heading['id']
            current_section_title = heading['title']

        has_math_xml = '<m:oMath' in paragraph._p.xml or '<m:oMathPara' in paragraph._p.xml
        if not has_math_xml and not _looks_like_formula_paragraph(text):
            continue

        previous_text = ' '.join(paragraphs[index - 2].text.split()) if index > 1 else ''
        next_paragraph_index, next_text = _next_nonempty_paragraph(paragraphs, index)
        formulas.append(
            {
                'paragraph_index': index,
                'text': text,
                'alignment': _alignment_name(paragraph.alignment),
                'has_math_xml': has_math_xml,
                'is_standalone': has_math_xml or _looks_like_formula_standalone(text),
                'equation_number': _extract_formula_number(text),
                'raw_equation_number': _extract_raw_formula_number(text),
                'prev_blank': not previous_text,
                'next_blank': not next_text,
                'prev_text': previous_text,
                'next_text': next_text,
                'next_paragraph_index': next_paragraph_index,
                'section_id': current_section_id,
                'section_title': current_section_title,
                'appendix_letter': _extract_appendix_letter(current_section_title),
            }
        )
    return formulas


def _starts_with_formula_explanation(text: str) -> bool:
    normalized = ' '.join((text or '').split()).lower()
    return normalized.startswith('где')


def _looks_like_formula_paragraph(text: str) -> bool:
    normalized = ' '.join((text or '').split())
    if not normalized:
        return False
    lowered = normalized.lower()
    if lowered.startswith('где'):
        return False
    if re.match(r'^[A-Za-zА-Яа-яЁё][^=]{0,40}[—–-]\s*.+$', normalized):
        return False
    if (_extract_formula_number(normalized) or _extract_raw_formula_number(normalized)) and any(symbol in normalized for symbol in ('=', '+', '-', '*', '/')):
        return True
    if '=' in normalized and re.search(r'[^\s=]{1,40}\s*=|=\s*[^\s=]', normalized):
        return True
    if re.search(r'[A-Za-z]\s*[+\-*/=]', normalized):
        return True
    if re.search(r'[0-9A-Za-z)\]]\s*[+\-*/]\s*[0-9A-Za-z(\[]', normalized):
        return True
    return False


def _looks_like_formula_standalone(text: str) -> bool:
    normalized = ' '.join((text or '').split())
    if not normalized:
        return False
    if len(normalized) > 160:
        return False
    if '=' not in normalized and not re.search(r'[+\-*/^]', normalized):
        return False
    before, formula, _ = _split_formula_sentence(normalized)
    if formula and before:
        return False
    if not formula and re.search(r'(?i)\b\u0444\u043e\u0440\u043c\u0443\u043b', normalized):
        return False
    return not bool(re.search(r'[.!?][^)]*$', normalized))


def _split_formula_sentence(text: str) -> tuple[str, str, str]:
    normalized = ' '.join((text or '').split())
    if not normalized or '=' not in normalized:
        return '', '', ''
    equal_index = normalized.find('=')
    start = _formula_split_start(normalized, equal_index)
    if start is None:
        return '', '', ''

    suffix = normalized[equal_index + 1:]
    stop_match = re.search(r'\s\u0433\u0434\u0435\b|[.;!?](?:\s|$)', suffix, flags=re.IGNORECASE)
    end = equal_index + 1 + (stop_match.start() if stop_match else len(suffix))

    formula_text = ' '.join(normalized[start:end].split()).rstrip(' .;!?')
    before = ' '.join(normalized[:start].split()).rstrip(' .;!?')
    after = ' '.join(normalized[end:].split()).lstrip(' .;!?')
    if not formula_text or '=' not in formula_text:
        return '', '', ''
    return before, formula_text, after


def _formula_split_start(text: str, equal_index: int) -> Optional[int]:
    prefix = text[:equal_index]
    lower_prefix = prefix.lower()
    markers = [
        'по формуле:',
        'по формуле',
        'в формуле:',
        'в формуле',
        'уравнение:',
        'формула:',
    ]
    marker_positions = [lower_prefix.rfind(marker) for marker in markers if lower_prefix.rfind(marker) >= 0]
    if marker_positions:
        marker_pos = max(marker_positions)
        for marker in markers:
            if lower_prefix.startswith(marker, marker_pos):
                return marker_pos + len(marker)

    match = re.search(r'([^\s=]{1,40}(?:\([^)]*\))?)\s*$', prefix)
    if not match:
        return None
    return match.start(1)


def _extract_formula_number(text: str) -> Optional[str]:
    match = re.search(r'\(((?:\d+(?:\.\d+)*)|(?:[A-ZА-ЯЁ]\.\d+))\)\s*$', ' '.join((text or '').split()))
    return match.group(1) if match else None


def _extract_raw_formula_number(text: str) -> Optional[str]:
    normalized = ' '.join((text or '').split())
    if _extract_formula_number(normalized):
        return None
    match = re.search(r'\s((?:\d+(?:\.\d+)*)|(?:[A-ZА-ЯЁ]\.\d+))\s*$', normalized)
    return match.group(1) if match else None


def _next_nonempty_paragraph(paragraphs: List[Any], index: int) -> tuple[Optional[int], str]:
    for next_index in range(index, len(paragraphs)):
        text = ' '.join(paragraphs[next_index].text.split())
        if text:
            return next_index + 1, text
    return None, ''

def _build_paragraph_meta_entry(paragraph: Any, index: int, text: str, alignment: str, style_name: str) -> Dict[str, Any]:
    return {
        'paragraph_index': index,
        'text': text,
        'alignment': alignment,
        'style': style_name,
        'first_line_indent_mm': _length_to_mm(_effective_first_line_indent(paragraph)),
        'line_spacing': _effective_line_spacing(paragraph),
        'font_size_pt_min': _paragraph_min_font_size_pt(paragraph),
        'font_family': _paragraph_font_family(paragraph),
        'has_bold_text': _paragraph_has_bold_text(paragraph),
        'has_non_black_text': _paragraph_has_non_black_text(paragraph),
    }


def _length_to_mm(length: Any) -> Optional[float]:
    if length is None:
        return None
    try:
        return round(float(length.mm), 1)
    except AttributeError:
        try:
            return round(float(length) / 36000.0, 1)
        except (TypeError, ValueError):
            return None


def _effective_first_line_indent(paragraph: Any) -> Any:
    indent = getattr(paragraph.paragraph_format, 'first_line_indent', None)
    if indent is not None:
        return indent
    return _resolve_style_paragraph_attr(getattr(paragraph, 'style', None), 'first_line_indent')


def _effective_line_spacing(paragraph: Any) -> Optional[float]:
    spacing = getattr(paragraph.paragraph_format, 'line_spacing', None)
    if spacing is None:
        spacing = _resolve_style_paragraph_attr(getattr(paragraph, 'style', None), 'line_spacing')
    if spacing is None:
        return None
    if isinstance(spacing, (int, float)):
        return round(float(spacing), 2)
    try:
        return round(float(spacing), 2)
    except (TypeError, ValueError):
        return None


def _paragraph_min_font_size_pt(paragraph: Any) -> Optional[float]:
    sizes: List[float] = []
    for run in getattr(paragraph, 'runs', []):
        if not ''.join(getattr(run, 'text', '').split()):
            continue
        size = _effective_run_font_size_pt(run, paragraph)
        if size is not None:
            sizes.append(size)
    return round(min(sizes), 1) if sizes else None


def _paragraph_font_family(paragraph: Any) -> Optional[str]:
    for run in getattr(paragraph, 'runs', []):
        if not ''.join(getattr(run, 'text', '').split()):
            continue
        family = _effective_run_font_name(run, paragraph)
        if family:
            return family
    return _resolve_style_font_attr(getattr(paragraph, 'style', None), 'name')


def _effective_run_font_name(run: Any, paragraph: Any) -> Optional[str]:
    name = getattr(getattr(run, 'font', None), 'name', None)
    if not name and getattr(run, 'style', None) is not None:
        name = _resolve_style_font_attr(run.style, 'name')
    if not name:
        name = _resolve_style_font_attr(getattr(paragraph, 'style', None), 'name')
    return str(name).strip() if name else None


def _paragraph_has_bold_text(paragraph: Any) -> bool:
    for run in getattr(paragraph, 'runs', []):
        if not ''.join(getattr(run, 'text', '').split()):
            continue
        if _effective_run_bold(run, paragraph):
            return True
    return False


def _paragraph_has_non_black_text(paragraph: Any) -> bool:
    for run in getattr(paragraph, 'runs', []):
        if not ''.join(getattr(run, 'text', '').split()):
            continue
        color = _effective_run_color_rgb(run, paragraph)
        if color and color not in {'000000', 'AUTO'}:
            return True
    return False


def _effective_run_font_size_pt(run: Any, paragraph: Any) -> Optional[float]:
    size = getattr(getattr(run, 'font', None), 'size', None)
    if size is None and getattr(run, 'style', None) is not None:
        size = _resolve_style_font_attr(run.style, 'size')
    if size is None:
        size = _resolve_style_font_attr(getattr(paragraph, 'style', None), 'size')
    if size is None:
        return None
    try:
        return round(float(size.pt), 1)
    except AttributeError:
        return None


def _effective_run_bold(run: Any, paragraph: Any) -> bool:
    bold = getattr(getattr(run, 'font', None), 'bold', None)
    if bold is None and getattr(run, 'style', None) is not None:
        bold = _resolve_style_font_attr(run.style, 'bold')
    if bold is None:
        bold = _resolve_style_font_attr(getattr(paragraph, 'style', None), 'bold')
    return bool(bold)


def _effective_run_color_rgb(run: Any, paragraph: Any) -> Optional[str]:
    color = getattr(getattr(getattr(run, 'font', None), 'color', None), 'rgb', None)
    if color is None and getattr(run, 'style', None) is not None:
        color = _resolve_style_font_color(run.style)
    if color is None:
        color = _resolve_style_font_color(getattr(paragraph, 'style', None))
    return str(color).upper() if color is not None else None


def _resolve_style_paragraph_attr(style: Any, attr: str) -> Any:
    current = style
    while current is not None:
        value = getattr(getattr(current, 'paragraph_format', None), attr, None)
        if value is not None:
            return value
        current = getattr(current, 'base_style', None)
    return None


def _resolve_style_font_attr(style: Any, attr: str) -> Any:
    current = style
    while current is not None:
        value = getattr(getattr(current, 'font', None), attr, None)
        if value is not None:
            return value
        current = getattr(current, 'base_style', None)
    return None


def _resolve_style_font_color(style: Any) -> Any:
    current = style
    while current is not None:
        color = getattr(getattr(getattr(current, 'font', None), 'color', None), 'rgb', None)
        if color is not None:
            return color
        current = getattr(current, 'base_style', None)
    return None


def _alignment_name(value: Any) -> str:
    if value is None:
        return 'unspecified'
    mapping = {
        0: 'left',
        1: 'center',
        2: 'right',
        3: 'justify',
        4: 'distribute',
    }
    try:
        return mapping.get(int(value), str(value).lower())
    except (TypeError, ValueError):
        return str(value).lower()


def _build_document_id(filename: str) -> str:
    stem = Path(filename).stem.lower()
    slug = re.sub(r"[^a-z0-9\u0430-\u044f\u0451]+", "_", stem, flags=re.IGNORECASE).strip('_')
    return slug or 'document'
