from io import BytesIO
import re
from typing import Iterable, List, Optional

from docx.shared import Mm, Pt, RGBColor

from app.fixing.document_fixer import apply_fixes
from app.schemas.document import DocumentInput
from app.schemas.issue import Issue
from app.utils.text import normalize_whitespace


def apply_fixes_to_source_docx(file_bytes: bytes, document: DocumentInput, issues: List[Issue]) -> bytes:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    source = Document(BytesIO(file_bytes))
    fixed = apply_fixes(document, issues)

    old_sections = {section.id: section for section in document.sections}
    new_sections = {section.id: section for section in fixed.sections}
    old_paragraphs = {paragraph.id: paragraph for paragraph in document.paragraphs}

    # 1. Apply explicit before/after fixes to source paragraphs.
    for issue in issues:
        suggestion = issue.suggestion
        if not hasattr(suggestion, 'before') or not hasattr(suggestion, 'after'):
            continue
        target_text = None
        if issue.location.paragraph_id and issue.location.paragraph_id in old_paragraphs:
            target_text = old_paragraphs[issue.location.paragraph_id].text
        elif suggestion.before:
            target_text = suggestion.before
        if target_text:
            paragraph = _find_paragraph_by_text(source, target_text)
            if paragraph is not None and not _paragraph_contains_drawing(paragraph):
                _replace_paragraph_text(paragraph, _replace_first(paragraph.text, suggestion.before, suggestion.after))

    # 2. Update section headings according to fixed parsed structure.
    existing_section_ids = [section.id for section in document.sections]
    for section_id in existing_section_ids:
        old_section = old_sections[section_id]
        new_section = new_sections.get(section_id)
        if new_section is None:
            continue
        old_heading = _format_heading(old_section)
        new_heading = _format_heading(new_section)
        if old_heading != new_heading:
            paragraph = _find_section_heading_paragraph(source, document, section_id)
            if paragraph is None:
                paragraph = _find_paragraph_by_text(source, old_heading)
            if paragraph is not None and not _paragraph_contains_drawing(paragraph):
                _replace_paragraph_text(paragraph, new_heading)

    # 2b. Apply paragraph-level layout and typography fixes before text replacements.
    _apply_alignment_fixes(source, document, issues)
    _apply_heading_layout_fixes(source, document, fixed, issues)
    _apply_typography_and_margin_fixes(source, document, issues)
    _apply_enumeration_fixes(source, document, issues)
    _apply_page_numbering_fixes(source, issues)

    # 3. Update captions that already exist in the document.
    for old_item, new_item in zip(document.figures, fixed.figures):
        _apply_caption_change(source, old_item.caption, new_item.caption, old_item.position.paragraph_index)
    for old_item, new_item in zip(document.tables, fixed.tables):
        _apply_caption_change(source, old_item.caption, new_item.caption, old_item.position.paragraph_index)

    _apply_table_structure_fixes(source, document, issues)
    _apply_formula_fixes(source, document, issues)
    _normalize_where_explanations(source)

    # 4. Append missing captions that were synthesized.
    for caption in _missing_captions(document, fixed):
        if not _paragraph_exists(source, caption):
            source.add_paragraph(caption)

    # 5. Append new sections that appeared only after fixes.
    old_ids = set(old_sections)
    for section in fixed.sections:
        if section.id in old_ids:
            continue
        source.add_paragraph(_format_heading(section))
        if section.text:
            for block in [part.strip() for part in section.text.split('\n\n') if part.strip()]:
                source.add_paragraph(block)

    output = BytesIO()
    source.save(output)
    return output.getvalue()


def _find_section_heading_paragraph(source_doc, parsed_document: DocumentInput, section_id: str):
    section_headings = parsed_document.meta.extras.get('section_headings', {}) if parsed_document.meta and parsed_document.meta.extras else {}
    payload = section_headings.get(section_id) or {}
    paragraph_index = payload.get('paragraph_index')
    paragraph = _find_docx_paragraph_by_index(source_doc, paragraph_index)
    if paragraph is not None:
        return paragraph
    heading_text = payload.get('text')
    if heading_text:
        return _find_paragraph_by_text(source_doc, heading_text)
    return None


def _apply_caption_change(source_doc, old_caption: str, new_caption: str, paragraph_index: Optional[int]) -> None:
    if not old_caption or not new_caption or old_caption == new_caption:
        return

    paragraph = None
    if paragraph_index is not None:
        candidate = _find_docx_paragraph_by_index(source_doc, paragraph_index)
        if candidate is not None and normalize_whitespace(candidate.text) == normalize_whitespace(old_caption):
            paragraph = candidate
    if paragraph is None:
        paragraph = _find_paragraph_by_text(source_doc, old_caption)
    if paragraph is None or _paragraph_contains_drawing(paragraph):
        return
    _replace_paragraph_text(paragraph, new_caption)


def _apply_table_structure_fixes(source_doc, parsed_document: DocumentInput, issues: List[Issue]) -> None:
    needs_move = any(issue.subtype == 'table_caption_below_table' for issue in issues)
    needs_header_cleanup = any(issue.subtype == 'invalid_table_header_punctuation' for issue in issues)
    needs_header_alignment = any(issue.subtype == 'table_header_alignment' for issue in issues)
    if not (needs_move or needs_header_cleanup or needs_header_alignment):
        return

    tables_meta = parsed_document.meta.extras.get('docx_tables_meta', []) if parsed_document.meta and parsed_document.meta.extras else []
    for issue in issues:
        if not issue.location.paragraph_id or not issue.location.paragraph_id.startswith('tbl_'):
            continue
        try:
            table_position = int(issue.location.paragraph_id.split('_', 1)[1]) - 1
        except ValueError:
            continue
        if table_position < 0 or table_position >= len(tables_meta) or table_position >= len(getattr(source_doc, 'tables', [])):
            continue
        meta = tables_meta[table_position]
        source_table = source_doc.tables[table_position]
        if issue.subtype == 'table_caption_below_table':
            caption_paragraph = _find_docx_paragraph_by_index(source_doc, meta.get('caption_paragraph_index'))
            if caption_paragraph is not None and not _paragraph_contains_drawing(caption_paragraph):
                source_table._tbl.addprevious(caption_paragraph._p)
        elif issue.subtype == 'invalid_table_header_punctuation' and source_table.rows:
            for cell in source_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    cleaned = paragraph.text.rstrip().rstrip('.')
                    if cleaned != paragraph.text:
                        _replace_paragraph_text(paragraph, cleaned)
        elif issue.subtype == 'table_header_alignment' and source_table.rows:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            for cell in source_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_formula_fixes(source_doc, parsed_document: DocumentInput, issues: List[Issue]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    formula_meta = parsed_document.meta.extras.get('docx_formulas', []) if parsed_document.meta and parsed_document.meta.extras else []
    if not formula_meta:
        return

    relevant_subtypes = {
        'formula_not_standalone',
        'formula_not_centered',
        'missing_formula_number',
        'formula_number_format_error',
        'formula_numbering_error',
        'appendix_formula_numbering_error',
        'invalid_formula_reference_format',
        'missing_formula_reference',
        'formula_where_colon',
        'formula_explanation_format_error',
        'formula_break_invalid',
    }
    issue_map = {}
    for issue in issues:
        if issue.subtype not in relevant_subtypes:
            continue
        paragraph_index = getattr(issue.location, 'paragraph_index', None)
        if paragraph_index is None:
            continue
        issue_map.setdefault(paragraph_index, set()).add(issue.subtype)

    if not issue_map:
        return

    expected_numbers = _expected_formula_numbers(formula_meta)

    for item in sorted(formula_meta, key=lambda payload: int(payload.get('paragraph_index') or 0), reverse=True):
        paragraph_index = item.get('paragraph_index')
        subtypes = issue_map.get(paragraph_index)
        if not subtypes:
            continue

        paragraph = _find_docx_paragraph_by_index(source_doc, paragraph_index)
        if paragraph is None:
            continue
        if _starts_with_where_text(paragraph.text):
            continue

        expected_number = expected_numbers.get(paragraph_index)
        has_math_xml = bool(item.get('has_math_xml'))
        formula_paragraph = paragraph
        before, candidate_formula, _ = _split_inline_formula_text(paragraph.text)
        moved_from_inline = (not has_math_xml) and (('formula_not_standalone' in subtypes) or bool(before and candidate_formula))
        if moved_from_inline:
            formula_paragraph = _move_inline_formula_to_separate_paragraph(source_doc, paragraph, expected_number)

        if 'formula_not_centered' in subtypes or moved_from_inline:
            formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        numbering_subtypes = {
            'missing_formula_number',
            'formula_number_format_error',
            'formula_numbering_error',
            'appendix_formula_numbering_error',
        }
        if (not has_math_xml) and expected_number and subtypes.intersection(numbering_subtypes):
            normalized_text = _normalize_formula_paragraph_text(formula_paragraph.text, expected_number)
            if normalized_text != formula_paragraph.text:
                _replace_paragraph_text(formula_paragraph, normalized_text)

        next_paragraph = _next_nonempty_paragraph_in_doc(source_doc, formula_paragraph)
        if next_paragraph is not None and _starts_with_where_text(next_paragraph.text):
            cleaned = normalize_whitespace(next_paragraph.text)
            if cleaned.lower().startswith('где:'):
                _replace_paragraph_text(next_paragraph, f"где {cleaned[4:].lstrip()}".strip())
        elif next_paragraph is not None and 'formula_explanation_format_error' in subtypes:
            normalized_next = normalize_whitespace(next_paragraph.text)
            if normalized_next and not _starts_with_where_text(normalized_next):
                _replace_paragraph_text(next_paragraph, f'где {normalized_next}')

        if next_paragraph is not None and 'formula_break_invalid' in subtypes:
            match = re.search(r'([=+\-*/])\s*$', normalize_whitespace(formula_paragraph.text))
            if match:
                sign = match.group(1)
                normalized_next = normalize_whitespace(next_paragraph.text)
                if normalized_next and not normalized_next.startswith(sign):
                    _replace_paragraph_text(next_paragraph, f'{sign} {normalized_next}')

        should_fix_reference = (not has_math_xml) and expected_number and subtypes.intersection({
            'invalid_formula_reference_format',
            'missing_formula_reference',
            'formula_number_format_error',
        })
        if should_fix_reference:
            _ensure_formula_reference_before(source_doc, formula_paragraph, expected_number)


def _starts_with_where_text(text: str) -> bool:
    normalized = normalize_whitespace(text).lower()
    return normalized.startswith('где')


def _move_inline_formula_to_separate_paragraph(source_doc, paragraph, expected_number: Optional[str]):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    before, formula_text, after = _split_inline_formula_text(paragraph.text)
    if not formula_text:
        return paragraph

    if expected_number:
        formula_text = _normalize_formula_paragraph_text(formula_text, expected_number)
        before = _normalize_formula_reference_in_text(before, expected_number)
    else:
        formula_text = normalize_whitespace(formula_text)

    if before:
        _replace_paragraph_text(paragraph, before)
        formula_paragraph = _insert_paragraph_after(paragraph, formula_text)
    else:
        _replace_paragraph_text(paragraph, formula_text)
        formula_paragraph = paragraph

    formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if after:
        _insert_paragraph_after(formula_paragraph, after)

    next_paragraph = _next_nonempty_paragraph_in_doc(source_doc, formula_paragraph)
    if next_paragraph is not None and _starts_with_where_text(next_paragraph.text):
        cleaned = normalize_whitespace(next_paragraph.text)
        if cleaned.lower().startswith('где:'):
            _replace_paragraph_text(next_paragraph, f"где {cleaned[4:].lstrip()}".strip())
    return formula_paragraph


def _normalize_formula_reference_in_text(text: str, number: str) -> str:
    normalized = normalize_whitespace(text)
    if not normalized:
        return normalized

    replaced = re.sub(r'(?i)(по\s+формуле)\s*:\s*$', lambda match: f"{match.group(1)} ({number})", normalized, count=1)
    if replaced != normalized:
        return replaced

    replaced = re.sub(r'(?i)(по\s+формуле)\s*$', lambda match: f"{match.group(1)} ({number})", normalized, count=1)
    if replaced != normalized:
        return replaced

    return normalized


def _split_inline_formula_text(text: str):
    normalized = normalize_whitespace(text)
    if not normalized or '=' not in normalized:
        return '', '', ''

    equal_index = normalized.find('=')
    start = _formula_split_start(normalized, equal_index)
    if start is None:
        return '', '', ''

    suffix = normalized[equal_index + 1:]
    stop_match = re.search(r'\s\u0433\u0434\u0435\b|[.;!?](?:\s|$)', suffix, flags=re.IGNORECASE)
    end = equal_index + 1 + (stop_match.start() if stop_match else len(suffix))

    formula_text = normalize_whitespace(normalized[start:end]).rstrip(' .;!?')
    before = normalize_whitespace(normalized[:start]).rstrip(' .;!?')
    after = normalize_whitespace(normalized[end:]).lstrip(' .;!?')
    if not formula_text or '=' not in formula_text:
        return '', '', ''
    return before, formula_text, after


def _formula_split_start(text: str, equal_index: int) -> Optional[int]:
    prefix = text[:equal_index]
    lower_prefix = prefix.lower()
    markers = [
        'по формуле:',
        'по формуле',
        'в формуле:',
        'в формуле',
        'уравнение:',
        'формула:',
    ]
    marker_positions = [lower_prefix.rfind(marker) for marker in markers if lower_prefix.rfind(marker) >= 0]
    if marker_positions:
        marker_pos = max(marker_positions)
        for marker in markers:
            if lower_prefix.startswith(marker, marker_pos):
                return marker_pos + len(marker)

    match = re.search(r'([^\s=]{1,40}(?:\([^)]*\))?)\s*$', prefix)
    if not match:
        return None
    return match.start(1)


def _expected_formula_numbers(formula_meta: List[dict]) -> dict[int, str]:
    result = {}
    simple_counter = 0
    appendix_counters = {}
    for item in formula_meta:
        paragraph_index = item.get('paragraph_index')
        if paragraph_index is None:
            continue
        appendix_letter = str(item.get('appendix_letter') or '').strip().upper()
        if appendix_letter:
            appendix_counters.setdefault(appendix_letter, 0)
            appendix_counters[appendix_letter] += 1
            result[paragraph_index] = f'{appendix_letter}.{appendix_counters[appendix_letter]}'
        else:
            simple_counter += 1
            result[paragraph_index] = str(simple_counter)
    return result


def _normalize_formula_paragraph_text(text: str, number: str) -> str:
    normalized = normalize_whitespace(text)
    normalized = re.sub(r'\s*\(((?:\d+(?:\.\d+)*)|(?:[A-ZА-ЯЁ]\.\d+))\)\s*$', '', normalized)
    normalized = re.sub(r'\s+((?:\d+(?:\.\d+)*)|(?:[A-ZА-ЯЁ]\.\d+))\s*$', '', normalized)
    normalized = normalized.rstrip()
    return f'{normalized} ({number})' if normalized else f'({number})'


def _ensure_formula_reference_before(source_doc, formula_paragraph, number: str) -> None:
    previous = _previous_nonempty_paragraph(source_doc, formula_paragraph)
    proper_reference = f'по формуле ({number})'
    if previous is None:
        _insert_paragraph_before(formula_paragraph, f'Расчет выполняют {proper_reference}.')
        return

    current_text = normalize_whitespace(previous.text)
    if re.search(rf'(?i)(?:в|по)?\s*формул(?:е|ы|у)?\s*\({re.escape(number)}\)', current_text):
        return

    if '=' in current_text or _starts_with_where_text(current_text):
        _insert_paragraph_before(formula_paragraph, f'Расчет выполняют {proper_reference}.')
        return

    replaced = re.sub(
        rf'(?i)((?:\u0432|\u043f\u043e)?\s*\u0444\u043e\u0440\u043c\u0443\u043b(?:\u0435|\u044b|\u0443)?\s+){re.escape(number)}(?=[\s.,;:!?]|$)',
        lambda match: f"{match.group(1).strip()} ({number})",
        current_text,
        count=1,
    )
    if replaced != current_text:
        _replace_paragraph_text(previous, replaced)
        return

    replaced = re.sub(
        r'(?i)(по\s+формуле)\s*:\s*$',
        lambda match: f"{match.group(1)} ({number})",
        current_text,
        count=1,
    )
    if replaced != current_text:
        _replace_paragraph_text(previous, replaced)
        return

    replaced = re.sub(
        r'(?i)(по\s+формуле)\s*$',
        lambda match: f"{match.group(1)} ({number})",
        current_text,
        count=1,
    )
    if replaced != current_text:
        _replace_paragraph_text(previous, replaced)
        return

    _insert_paragraph_before(formula_paragraph, f'Расчет выполняют {proper_reference}.')


def _previous_nonempty_paragraph(document, paragraph):
    paragraphs = getattr(document, 'paragraphs', None) or []
    for index, item in enumerate(paragraphs):
        if item._p is not paragraph._p:
            continue
        for previous_index in range(index - 1, -1, -1):
            candidate = paragraphs[previous_index]
            if normalize_whitespace(candidate.text):
                return candidate
        return None
    return None


def _insert_paragraph_before(paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _normalize_where_explanations(source_doc) -> None:
    for paragraph in getattr(source_doc, 'paragraphs', []) or []:
        normalized = normalize_whitespace(paragraph.text)
        if normalized.lower().startswith('где:'):
            tail = normalized[4:].lstrip()
            _replace_paragraph_text(paragraph, f'где {tail}'.strip())


def _find_docx_paragraph_by_index(document, paragraph_index: Optional[int]):
    if paragraph_index is None or paragraph_index <= 0:
        return None
    paragraphs = getattr(document, 'paragraphs', None) or []
    if paragraph_index > len(paragraphs):
        return None
    return paragraphs[paragraph_index - 1]


def _next_nonempty_paragraph_in_doc(document, paragraph):
    paragraphs = getattr(document, 'paragraphs', None) or []
    for index, item in enumerate(paragraphs):
        if item._p is not paragraph._p:
            continue
        for next_index in range(index + 1, len(paragraphs)):
            candidate = paragraphs[next_index]
            if normalize_whitespace(candidate.text):
                return candidate
        return None
    return None


def _missing_captions(original: DocumentInput, fixed: DocumentInput) -> List[str]:
    captions: List[str] = []
    for old_item, new_item in zip(original.figures, fixed.figures):
        if not old_item.caption and new_item.caption:
            captions.append(new_item.caption)
    for old_item, new_item in zip(original.tables, fixed.tables):
        if not old_item.caption and new_item.caption:
            captions.append(new_item.caption)
    return captions


def _apply_heading_layout_fixes(source_doc, parsed_document: DocumentInput, fixed_document: DocumentInput, issues: List[Issue]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    issue_types = {issue.subtype for issue in issues}
    relevant_issue_types = {
        'main_heading_centered',
        'main_heading_indent_invalid',
        'structural_heading_numbered',
        'structural_heading_not_uppercase',
        'structural_heading_not_centered',
        'structural_heading_has_indent',
        'structural_heading_invalid_font_color',
        'appendix_heading_not_centered',
        'appendix_heading_single_line',
        'heading_invalid_font_color',
    }
    if not issue_types.intersection(relevant_issue_types):
        return

    structural_titles = {
        'содержание',
        'оглавление',
        'реферат',
        'список использованных источников',
        'список литературы',
        'библиографический список',
        'источники и литература',
    }
    fixed_sections = {section.id: section for section in fixed_document.sections}
    for section in parsed_document.sections:
        paragraph = _find_section_heading_paragraph(source_doc, parsed_document, section.id)
        if paragraph is None:
            continue
        fixed_section = fixed_sections.get(section.id, section)
        normalized = normalize_whitespace(fixed_section.title).lower()
        is_appendix = normalized.startswith('приложение')
        is_structural = normalized in structural_titles

        if is_appendix:
            letter, tail = _split_appendix_heading(fixed_section.title)
            appendix_line = f'ПРИЛОЖЕНИЕ {letter}' if letter else 'ПРИЛОЖЕНИЕ'
            _replace_paragraph_text(paragraph, appendix_line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.left_indent = Mm(0)
            paragraph.paragraph_format.right_indent = Mm(0)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
            if tail:
                title_paragraph = _next_paragraph(source_doc, paragraph)
                if title_paragraph is None or _looks_like_heading_text(title_paragraph.text) or _looks_like_caption_text(title_paragraph.text):
                    title_paragraph = _insert_paragraph_after(paragraph, tail)
                else:
                    _replace_paragraph_text(title_paragraph, tail)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_paragraph.paragraph_format.first_line_indent = Mm(0)
                title_paragraph.paragraph_format.left_indent = Mm(0)
                title_paragraph.paragraph_format.right_indent = Mm(0)
                for run in title_paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.bold = True
            continue

        if is_structural:
            _replace_paragraph_text(paragraph, normalize_whitespace(fixed_section.title).upper())
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Mm(0)
            paragraph.paragraph_format.left_indent = Mm(0)
            paragraph.paragraph_format.right_indent = Mm(0)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Mm(12.5)
        paragraph.paragraph_format.left_indent = Mm(0)
        paragraph.paragraph_format.right_indent = Mm(0)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True


def _split_appendix_heading(title: str):
    normalized = normalize_whitespace(title)
    match = re.match(r'^Приложение\s+([\u0410-\u042fA-Z])(?:\s*[—–-]\s*(.+)|\s+(.+))?$', normalized, re.IGNORECASE)
    if not match:
        return None, normalized
    return match.group(1).upper(), normalize_whitespace(match.group(2) or match.group(3) or '')


def _next_paragraph(document, paragraph):
    paragraphs = getattr(document, 'paragraphs', None) or []
    for index, item in enumerate(paragraphs):
        if item._p is paragraph._p and index + 1 < len(paragraphs):
            return paragraphs[index + 1]
    return None


def _insert_paragraph_after(paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _looks_like_heading_text(text: str) -> bool:
    normalized = normalize_whitespace(text)
    if not normalized:
        return False
    return bool(re.match(r'^(\d+(?:\.\d+)*)\s+.+$', normalized) or normalized.upper().startswith('ПРИЛОЖЕНИЕ'))



def _looks_like_caption_text(text: str) -> bool:
    normalized = normalize_whitespace(text)
    return normalized.startswith('??????? ') or normalized.startswith('??????? ')

def _apply_typography_and_margin_fixes(source_doc, parsed_document: DocumentInput, issues: List[Issue]) -> None:
    page_size_fix_needed = any(issue.subtype == 'invalid_page_size' for issue in issues)
    if page_size_fix_needed:
        for section in getattr(source_doc, 'sections', []):
            width_mm = float(getattr(section.page_width, 'mm', 0) or 0)
            height_mm = float(getattr(section.page_height, 'mm', 0) or 0)
            if width_mm > height_mm:
                section.page_width = Mm(297)
                section.page_height = Mm(210)
            else:
                section.page_width = Mm(210)
                section.page_height = Mm(297)

    margin_fix_needed = any(issue.subtype == 'invalid_page_margins' for issue in issues)
    if margin_fix_needed:
        for section in getattr(source_doc, 'sections', []):
            section.left_margin = Mm(30)
            section.right_margin = Mm(15)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)

    typography_issue_types = {
        'invalid_first_line_indent',
        'invalid_line_spacing',
        'invalid_font_size',
        'invalid_font_family',
        'invalid_font_color',
        'unexpected_bold_text',
    }
    triggered_typography_fixes = {issue.subtype for issue in issues if issue.subtype in typography_issue_types}
    if not triggered_typography_fixes:
        return

    body_font_fix_types = {'invalid_font_size', 'invalid_font_family', 'invalid_font_color', 'unexpected_bold_text'}
    heading_font_fix_types = {'invalid_font_size', 'invalid_font_family', 'invalid_font_color', 'heading_invalid_font_color', 'structural_heading_invalid_font_color'}

    for paragraph in _collect_body_paragraphs_for_typography_fix(source_doc, parsed_document):
        if _paragraph_contains_drawing(paragraph):
            continue
        if 'invalid_first_line_indent' in triggered_typography_fixes:
            paragraph.paragraph_format.first_line_indent = Mm(12.5)
        if 'invalid_line_spacing' in triggered_typography_fixes:
            paragraph.paragraph_format.line_spacing = 1.5
        if any(subtype in triggered_typography_fixes for subtype in body_font_fix_types):
            for run in paragraph.runs:
                if not ''.join(run.text.split()):
                    continue
                if 'invalid_font_size' in triggered_typography_fixes:
                    run.font.size = Pt(12)
                if 'invalid_font_family' in triggered_typography_fixes:
                    run.font.name = 'Times New Roman'
                if 'invalid_font_color' in triggered_typography_fixes:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                if 'unexpected_bold_text' in triggered_typography_fixes:
                    run.font.bold = False

    for paragraph in _collect_heading_paragraphs_for_typography_fix(source_doc, parsed_document):
        if _paragraph_contains_drawing(paragraph):
            continue
        if any(subtype in triggered_typography_fixes for subtype in heading_font_fix_types):
            style = getattr(paragraph, 'style', None)
            if style is not None:
                if 'invalid_font_size' in triggered_typography_fixes:
                    style.font.size = Pt(12)
                if 'invalid_font_family' in triggered_typography_fixes:
                    style.font.name = 'Times New Roman'
                if 'invalid_font_color' in triggered_typography_fixes or 'heading_invalid_font_color' in triggered_typography_fixes:
                    style.font.color.rgb = RGBColor(0, 0, 0)
            for run in paragraph.runs:
                if not ''.join(run.text.split()):
                    continue
                if 'invalid_font_size' in triggered_typography_fixes:
                    run.font.size = Pt(12)
                if 'invalid_font_family' in triggered_typography_fixes:
                    run.font.name = 'Times New Roman'
                if 'invalid_font_color' in triggered_typography_fixes or 'heading_invalid_font_color' in triggered_typography_fixes:
                    run.font.color.rgb = RGBColor(0, 0, 0)


def _find_paragraph_for_issue(source_doc, parsed_document: DocumentInput, issue: Issue):
    paragraph = _find_docx_paragraph_by_index(source_doc, getattr(issue.location, 'paragraph_index', None))
    if paragraph is not None:
        return paragraph

    if issue.location.paragraph_id:
        for payload in parsed_document.paragraphs:
            if payload.id == issue.location.paragraph_id and payload.position.paragraph_index is not None:
                paragraph = _find_docx_paragraph_by_index(source_doc, payload.position.paragraph_index)
                if paragraph is not None:
                    return paragraph
                break

    target_text = _extract_quoted_fragment(issue.evidence)
    if target_text:
        paragraph = _find_paragraph_by_text(source_doc, target_text)
        if paragraph is not None:
            return paragraph

    for payload in parsed_document.paragraphs:
        if payload.id == issue.location.paragraph_id and payload.text:
            return _find_paragraph_by_text(source_doc, payload.text)
    return None


def _collect_body_paragraphs_for_typography_fix(source_doc, parsed_document: DocumentInput):
    paragraphs = []
    seen_indexes = set()
    heading_texts = {_format_heading(section) for section in parsed_document.sections}
    for payload in parsed_document.paragraphs:
        paragraph_index = payload.position.paragraph_index
        if paragraph_index is None or paragraph_index in seen_indexes:
            continue
        text = ' '.join((payload.text or '').split())
        if not text:
            continue
        if text in heading_texts:
            continue
        paragraph = _find_docx_paragraph_by_index(source_doc, paragraph_index)
        if paragraph is None:
            continue
        paragraphs.append(paragraph)
        seen_indexes.add(paragraph_index)
    return paragraphs


def _collect_heading_paragraphs_for_typography_fix(source_doc, parsed_document: DocumentInput):
    paragraphs = []
    seen_indexes = set()
    section_headings = parsed_document.meta.extras.get('section_headings', {}) if parsed_document.meta and parsed_document.meta.extras else {}
    for section in parsed_document.sections:
        payload = section_headings.get(section.id, {})
        paragraph_index = payload.get('paragraph_index')
        paragraph = None
        if paragraph_index is not None and paragraph_index not in seen_indexes:
            paragraph = _find_docx_paragraph_by_index(source_doc, paragraph_index)
            if paragraph is not None:
                seen_indexes.add(paragraph_index)
        if paragraph is None:
            paragraph = _find_paragraph_by_text(source_doc, _format_heading(section))
        if paragraph is None:
            continue
        paragraphs.append(paragraph)
    return paragraphs


def _apply_enumeration_fixes(source_doc, parsed_document: DocumentInput, issues: List[Issue]) -> None:
    relevant = {'invalid_enumeration_marker', 'enumeration_indent_invalid'}
    if not any(issue.subtype in relevant for issue in issues):
        return

    for issue in issues:
        if issue.subtype not in relevant:
            continue
        paragraph = _find_paragraph_for_issue(source_doc, parsed_document, issue)
        if paragraph is None or _paragraph_contains_drawing(paragraph):
            continue
        text = normalize_whitespace(paragraph.text)
        if issue.subtype == 'invalid_enumeration_marker':
            normalized = _normalize_enumeration_marker(text)
            if normalized != text:
                _replace_paragraph_text(paragraph, normalized)
        if issue.subtype in {'invalid_enumeration_marker', 'enumeration_indent_invalid'}:
            paragraph.paragraph_format.first_line_indent = Mm(12.5)


def _apply_page_numbering_fixes(source_doc, issues: List[Issue]) -> None:
    relevant_subtypes = {
        'missing_page_numbering',
        'page_number_not_centered',
        'title_page_number_visible',
        'page_numbering_restart',
    }
    if not any(issue.subtype in relevant_subtypes for issue in issues):
        return

    for index, section in enumerate(getattr(source_doc, 'sections', []), start=1):
        footer = getattr(section, 'footer', None)
        first_footer = getattr(section, 'first_page_footer', None)
        header = getattr(section, 'header', None)
        first_header = getattr(section, 'first_page_header', None)
        even_header = getattr(section, 'even_page_header', None)

        for region in (footer, first_footer, header, first_header, even_header):
            if region is None:
                continue
            try:
                region.is_linked_to_previous = False
            except Exception:
                pass

        if index == 1:
            try:
                section.different_first_page_header_footer = True
            except Exception:
                pass
            if first_footer is not None:
                _remove_footer_page_fields(first_footer)
            if first_header is not None:
                _remove_footer_page_fields(first_header)

        if header is not None:
            _remove_footer_page_fields(header)
        if even_header is not None:
            _remove_footer_page_fields(even_header)

        _remove_page_number_restart(section)
        if footer is not None:
            _ensure_centered_footer_page_field(footer)


def _remove_footer_page_fields(footer) -> None:
    for paragraph in getattr(footer, 'paragraphs', []) or []:
        if _paragraph_has_page_field(paragraph):
            _replace_paragraph_text(paragraph, '')
            paragraph.alignment = None


def _ensure_centered_footer_page_field(footer) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = _find_footer_page_paragraph(footer)
    if paragraph is None:
        paragraphs = list(getattr(footer, 'paragraphs', []) or [])
        paragraph = paragraphs[0] if paragraphs else footer.add_paragraph()
    for candidate in getattr(footer, 'paragraphs', []) or []:
        if candidate is paragraph:
            continue
        if _paragraph_has_page_field(candidate):
            _replace_paragraph_text(candidate, '')
    _replace_paragraph_with_page_field(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _find_footer_page_paragraph(footer):
    for paragraph in getattr(footer, 'paragraphs', []) or []:
        if _paragraph_has_page_field(paragraph):
            return paragraph
    return None


def _paragraph_has_page_field(paragraph) -> bool:
    xml = getattr(getattr(paragraph, '_p', None), 'xml', '') or ''
    upper_xml = xml.upper()
    return 'PAGE' in upper_xml and ('INSTRTEXT' in upper_xml or 'FLDSIMPLE' in upper_xml)


def _replace_paragraph_with_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = paragraph._p
    for child in list(p):
        if child.tag.endswith('}pPr'):
            continue
        p.remove(child)

    begin_run = paragraph.add_run()._r
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    begin_run.append(fld_begin)

    instr_run = paragraph.add_run()._r
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    instr_run.append(instr)

    separate_run = paragraph.add_run()._r
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    separate_run.append(fld_separate)

    text_run = paragraph.add_run()._r
    page_text = OxmlElement('w:t')
    page_text.text = '1'
    text_run.append(page_text)

    end_run = paragraph.add_run()._r
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    end_run.append(fld_end)


def _remove_page_number_restart(section) -> None:
    try:
        from docx.oxml.ns import qn
    except Exception:
        return
    sect_pr = getattr(section, '_sectPr', None)
    if sect_pr is None:
        return
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is not None:
        sect_pr.remove(pg_num_type)


def _apply_alignment_fixes(source_doc, parsed_document: DocumentInput, issues: List[Issue]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for issue in issues:
        target_alignment = None
        if issue.subtype == 'figure_caption_not_centered':
            target_alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif issue.subtype == 'title_page_right_alignment':
            target_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif issue.subtype == 'title_page_center_alignment':
            target_alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif issue.subtype == 'body_text_not_justified':
            target_alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif issue.subtype == 'table_header_alignment':
            target_alignment = None

        if target_alignment is None:
            continue

        paragraph = _find_docx_paragraph_by_index(source_doc, getattr(issue.location, 'paragraph_index', None))
        target_text = _extract_quoted_fragment(issue.evidence)
        if paragraph is None and target_text:
            paragraph = _find_paragraph_by_text(source_doc, target_text)
        if paragraph is None and issue.subtype == 'figure_caption_not_centered':
            paragraph = _find_paragraph_by_text(source_doc, _find_figure_caption_for_issue(parsed_document, issue))
        if paragraph is None or _paragraph_contains_drawing(paragraph):
            continue
        paragraph.alignment = target_alignment


def _extract_quoted_fragment(text: str) -> Optional[str]:
    match = re.search(r"'([^']+)'", text or '')
    if match:
        return ' '.join(match.group(1).split())
    return None


def _find_figure_caption_for_issue(document: DocumentInput, issue: Issue) -> Optional[str]:
    target_page = issue.location.page
    for figure in document.figures:
        if target_page is not None and figure.position.page == target_page and figure.caption:
            return figure.caption
    if document.figures:
        return document.figures[0].caption or None
    return None


def _normalize_enumeration_marker(text: str) -> str:
    normalized = normalize_whitespace(text)
    normalized = re.sub(r'^(?:[????])\s+', '- ', normalized)
    normalized = re.sub(r'^(\d+)\.\s+', r'\1) ', normalized)
    normalized = re.sub(r'^([A-Za-z?-??-???])\.\s+', r'\1) ', normalized)
    return normalized


def _replace_first(source: str, before: str, after: str) -> str:
    if before and before in source:
        return source.replace(before, after, 1)
    return after or source


def _format_heading(section) -> str:
    if section.number:
        return f'{section.number} {section.title}'.strip()
    return section.title.strip()


def _paragraph_exists(document, text: str) -> bool:
    return _find_paragraph_by_text(document, text) is not None


def _find_paragraph_by_text(document, text: str):
    target = ' '.join((text or '').split())
    if not target:
        return None
    for paragraph in _iter_paragraphs(document):
        current = ' '.join(paragraph.text.split())
        if current == target:
            return paragraph
    return None


def _iter_paragraphs(parent) -> Iterable:
    from docx.document import Document as DocumentType
    from docx.table import _Cell, Table

    if isinstance(parent, DocumentType):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from _iter_table_paragraphs(table)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _paragraph_contains_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath('.//*[local-name()="drawing" or local-name()="object" or local-name()="pict"]'))


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith('}pPr'):
            continue
        p.remove(child)
    paragraph.add_run(new_text)


